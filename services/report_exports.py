"""Professional, deterministic PDF and DOCX analysis exports."""

from __future__ import annotations

from io import BytesIO
from xml.sax.saxutils import escape


def build_pdf(report: str, title: str) -> bytes:
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer

    stream = BytesIO()
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("ReportTitle", parent=styles["Title"], alignment=TA_CENTER, textColor="#12506A")
    doc = SimpleDocTemplate(stream, pagesize=A4, rightMargin=20*mm, leftMargin=20*mm, topMargin=18*mm, bottomMargin=18*mm,
                            title=title, author="Veraxis")
    story = [Paragraph(escape(title), title_style), Spacer(1, 8*mm)]
    for line in report.splitlines():
        text = line.strip()
        if not text:
            story.append(Spacer(1, 3*mm)); continue
        if text.startswith("#"):
            level = min(3, len(text) - len(text.lstrip("#")))
            story.append(Paragraph(escape(text.lstrip("# ")), styles[f"Heading{level}"]))
        else:
            story.append(Paragraph(escape(text.lstrip("- ")), styles["BodyText"]))
    doc.build(story)
    return stream.getvalue()


def build_docx(report: str, title: str) -> bytes:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor

    document = Document()
    section = document.sections[0]
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"; normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    heading_tokens = {
        "Heading 1": (16, "2E74B5", 16, 8),
        "Heading 2": (13, "2E74B5", 12, 6),
        "Heading 3": (12, "1F4D78", 8, 4),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = document.styles[style_name]
        style.font.name = "Calibri"; style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before); style.paragraph_format.space_after = Pt(after)
    heading = document.add_heading(title, 0)
    heading.runs[0].font.name = "Calibri"; heading.runs[0].font.size = Pt(24)
    heading.runs[0].font.color.rgb = RGBColor(18, 80, 106)
    heading.paragraph_format.space_after = Pt(16)
    for line in report.splitlines():
        text = line.strip()
        if not text:
            continue
        if text.startswith("#"):
            level = min(3, len(text) - len(text.lstrip("#")))
            document.add_heading(text.lstrip("# "), level=level)
        elif text.startswith("- "):
            document.add_paragraph(text[2:], style="List Bullet")
        else:
            document.add_paragraph(text)
    stream = BytesIO(); document.save(stream); return stream.getvalue()
