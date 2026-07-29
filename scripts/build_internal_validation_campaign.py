"""Execute 20 reproducible cases and build the human-review DOCX/JSON package."""

from __future__ import annotations

import json
import sys
from datetime import datetime, timezone
from pathlib import Path

import pandas as pd

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from services.analysis_engine import AnalysisEngine
from validation_lab.functional_benchmark import benchmark_cases, _contains


def selected_cases() -> list[dict]:
    cases = benchmark_cases()
    baseline = cases[:30]
    adversarial = cases[30:]
    return baseline[:15] + adversarial[:5]


def execute_campaign() -> dict:
    engine = AnalysisEngine()
    results = []
    for index, case in enumerate(selected_cases(), start=1):
        payload = engine.run(case["prompt"], pd.DataFrame(case["records"]), source_type="internal_validation")
        actual = payload["deterministic_results"]
        actual_type = payload["analysis_plan"]["analysis_type"]
        expected_status = case.get("expected_status")
        automatic_pass = (
            actual_type == case["expected_type"]
            and (not expected_status or actual.get("status") == expected_status)
            and _contains(actual, case["expected"])
        )
        results.append({
            "review_id": f"REV-{index:03d}",
            "benchmark_id": case["id"],
            "domain": case["domain"],
            "category": case["expected_type"],
            "objective": objective_for(case["expected_type"]),
            "prompt": case["prompt"],
            "dataset": {
                "row_count": len(case["records"]),
                "columns": list(pd.DataFrame(case["records"]).columns),
                "records": case["records"],
            },
            "expected": case["expected"],
            "actual_type": actual_type,
            "actual": actual,
            "automatic_result": "passed" if automatic_pass else "failed",
            "human_feedback": {
                "outcome": None,
                "rating": None,
                "clarity": None,
                "usefulness": None,
                "notes": None,
            },
        })
    return {
        "campaign": "Veraxis internal single-user validation",
        "evidence_class": "internal_single_user_not_independent_beta_feedback",
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "total": len(results),
        "automatic_passed": sum(item["automatic_result"] == "passed" for item in results),
        "cases": results,
    }


def objective_for(analysis_type: str) -> str:
    return {
        "count_occurrences": "Verificare conteggi, categorie e totale dei record.",
        "numeric_aggregation": "Verificare aggregazione, metrica e raggruppamento.",
        "top_n": "Verificare ordinamento, limite e valore aggregato.",
        "null_detection": "Verificare completezza e conteggio dei valori mancanti.",
        "duplicate_detection": "Verificare individuazione e quantificazione dei duplicati.",
    }.get(analysis_type, "Verificare che il contratto analitico sia rispettato.")


def build_docx(campaign: dict, output: Path) -> None:
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    document = Document()
    section = document.sections[0]
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"; normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ):
        style = document.styles[name]
        style.font.name = "Calibri"; style.font.size = Pt(size); style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before); style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.text = "VERAXIS · CAMPAGNA DI VALIDAZIONE INTERNA"
    header.style = document.styles["Caption"]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Documento di revisione single-user · ")
    field = OxmlElement("w:fldSimple"); field.set(qn("w:instr"), "PAGE"); footer._p.append(field)

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("Campagna di validazione interna")
    run.bold = True; run.font.name = "Calibri"; run.font.size = Pt(24); run.font.color.rgb = RGBColor.from_string("12506A")
    subtitle = document.add_paragraph("20 casi eseguiti automaticamente e predisposti per il feedback umano")
    subtitle.style = document.styles["Subtitle"]
    document.add_heading("Come compilare il documento", level=1)
    for text in (
        "Confronta il risultato atteso con il risultato effettivo.",
        "Compila tutti i campi contrassegnati con [COMPILARE].",
        "Usa outcome corretto, parziale o errato e rating da 1 a 5.",
        "Non modificare gli ID REV: servono per importare correttamente il feedback.",
    ):
        document.add_paragraph(text, style="List Number")
    note = document.add_paragraph()
    note.add_run("Classificazione dell'evidenza: ").bold = True
    note.add_run("validazione interna single-user; non equivale a feedback beta indipendente.")
    document.add_heading("Riepilogo automatico", level=1)
    summary = document.add_table(rows=2, cols=3)
    set_table_geometry(summary, [3120, 3120, 3120])
    for cell, value in zip(summary.rows[0].cells, ("Casi", "Superati", "Da verificare")):
        style_cell(cell, value, header=True)
    for cell, value in zip(summary.rows[1].cells, (campaign["total"], campaign["automatic_passed"], campaign["total"])):
        style_cell(cell, str(value))

    for item in campaign["cases"]:
        document.add_page_break()
        document.add_heading(f"{item['review_id']} · {item['benchmark_id']}", level=1)
        meta = document.add_table(rows=4, cols=2)
        set_table_geometry(meta, [2700, 6660])
        for row, (label, value) in zip(meta.rows, (
            ("Dominio", item["domain"]), ("Categoria", item["category"]),
            ("Esito automatico", item["automatic_result"].upper()),
            ("Dataset", f"{item['dataset']['row_count']} righe · {', '.join(item['dataset']['columns'])}"),
        )):
            style_cell(row.cells[0], label, header=True); style_cell(row.cells[1], str(value))
        add_labeled_text(document, "Obiettivo", item["objective"])
        add_labeled_text(document, "Domanda sottoposta al sistema", item["prompt"])
        add_labeled_text(document, "Risultato atteso", compact_json(item["expected"]))
        add_labeled_text(document, "Risultato effettivo", compact_json(item["actual"]))
        document.add_heading("Valutazione umana", level=2)
        feedback = document.add_table(rows=5, cols=2)
        set_table_geometry(feedback, [2700, 6660])
        prompts = (
            ("Outcome", "[COMPILARE: corretto / parziale / errato]"),
            ("Rating", "[COMPILARE: 1-5]"),
            ("Chiarezza", "[COMPILARE: 1-5]"),
            ("Utilità", "[COMPILARE: 1-5]"),
            ("Note", "[COMPILARE: osservazioni, differenze o correzioni suggerite]"),
        )
        for row, (label, value) in zip(feedback.rows, prompts):
            style_cell(row.cells[0], label, header=True); style_cell(row.cells[1], value)

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def add_labeled_text(document, label: str, text: str) -> None:
    document.add_heading(label, level=2)
    paragraph = document.add_paragraph(text)
    paragraph.paragraph_format.keep_together = True


def compact_json(value) -> str:
    return json.dumps(value, ensure_ascii=False, indent=2, default=str)


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    table.alignment = 0
    table.autofit = False
    properties = table._tbl.tblPr
    width = properties.first_child_found_in("w:tblW")
    if width is None:
        width = OxmlElement("w:tblW"); properties.append(width)
    width.set(qn("w:w"), str(sum(widths_dxa))); width.set(qn("w:type"), "dxa")
    indent = OxmlElement("w:tblInd"); indent.set(qn("w:w"), "120"); indent.set(qn("w:type"), "dxa"); properties.append(indent)
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for value in widths_dxa:
        column = OxmlElement("w:gridCol"); column.set(qn("w:w"), str(value)); grid.append(column)
    for row in table.rows:
        for cell, value in zip(row.cells, widths_dxa):
            cell.width = value
            tcw = cell._tc.get_or_add_tcPr().get_or_add_tcW(); tcw.set(qn("w:w"), str(value)); tcw.set(qn("w:type"), "dxa")
            margins = OxmlElement("w:tcMar")
            for side, amount in (("top",80),("bottom",80),("start",120),("end",120)):
                node=OxmlElement(f"w:{side}"); node.set(qn("w:w"),str(amount)); node.set(qn("w:type"),"dxa"); margins.append(node)
            cell._tc.get_or_add_tcPr().append(margins)


def style_cell(cell, text: str, header: bool = False) -> None:
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]; paragraph.paragraph_format.space_after = 0
    run = paragraph.add_run(text); run.bold = header
    if header:
        shade = OxmlElement("w:shd"); shade.set(qn("w:fill"), "E8EEF5"); cell._tc.get_or_add_tcPr().append(shade)


def main() -> int:
    output_dir = ROOT / "validation_lab" / "deliverables"
    campaign = execute_campaign()
    json_path = output_dir / "Veraxis_Campagna_Validazione_Interna.json"
    docx_path = output_dir / "Veraxis_Campagna_Validazione_Interna.docx"
    output_dir.mkdir(parents=True, exist_ok=True)
    json_path.write_text(json.dumps(campaign, ensure_ascii=False, indent=2, default=str) + "\n", encoding="utf-8")
    build_docx(campaign, docx_path)
    print(json.dumps({"docx": str(docx_path), "json": str(json_path), "total": campaign["total"], "passed": campaign["automatic_passed"]}))
    return 0 if campaign["automatic_passed"] == campaign["total"] else 2


if __name__ == "__main__":
    raise SystemExit(main())
