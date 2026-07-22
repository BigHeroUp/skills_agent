"""Build the illustrated Veraxis product, administration and operations manual."""

from __future__ import annotations

from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from scripts.build_user_manual import (
    ASSETS as USER_ASSETS,
    BLUE,
    GOLD,
    GRAY,
    LIGHT_BLUE,
    LIGHT_GOLD,
    LIGHT_GREEN,
    LIGHT_RED,
    NAVY,
    TEAL,
    add_page_field,
    bullet,
    callout,
    configure,
    set_font,
    set_table_geometry,
    shade,
)


ROOT = Path(__file__).resolve().parents[1]
ASSETS = ROOT / "docs" / "product_manual_assets"
OUTPUT = ROOT / "docs" / "deliverables" / "Manuale_prodotto_Veraxis.docx"


def font(size, bold=False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/Library/Fonts/Arial Bold.ttf" if bold else "/Library/Fonts/Arial.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def arrow(draw, start, end, color="#4AAFC1", width=5):
    draw.line([start, end], fill=color, width=width)
    x, y = end
    draw.polygon([(x, y), (x - 14, y - 9), (x - 14, y + 9)], fill=color)


def box(draw, xy, title, subtitle, fill="#102A3A", outline="#54D7EA"):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=22, fill=fill, outline=outline, width=4)
    draw.text((x1 + 22, y1 + 18), title, fill="white", font=font(30, True))
    lines = subtitle.split("\n")
    for idx, line in enumerate(lines):
        draw.text((x1 + 22, y1 + 64 + idx * 30), line, fill="#C9E5EC", font=font(22))


def create_architecture_diagram():
    image = Image.new("RGB", (1600, 900), "#061018")
    draw = ImageDraw.Draw(image)
    draw.text((60, 35), "Architettura commerciale Veraxis", fill="white", font=font(42, True))
    boxes = [
        ((70, 180, 350, 350), "Utente / Client", "Portale web\nREST API"),
        ((440, 180, 720, 350), "Gateway", "Nginx + HTTPS\nlimiti e header"),
        ((810, 180, 1110, 350), "Platform API", "Auth, RBAC\njob e storico"),
        ((1210, 100, 1515, 270), "PostgreSQL", "tenant, utenti\nanalisi, feedback"),
        ((1210, 360, 1515, 530), "Redis / RQ", "coda durevole\nretry, cancel"),
        ((810, 580, 1110, 760), "Coordinator", "pipeline di\nproduzione"),
        ((440, 580, 720, 760), "Product Intelligence", "KG, Experience\nRecommendation, Decision"),
        ((70, 580, 350, 760), "Storage tenant", "Knowledge Graph\nExperience e history"),
    ]
    for xy, title, subtitle in boxes:
        box(draw, xy, title, subtitle)
    arrow(draw, (350, 265), (440, 265))
    arrow(draw, (720, 265), (810, 265))
    arrow(draw, (1110, 220), (1210, 185))
    arrow(draw, (1110, 310), (1210, 440))
    arrow(draw, (1360, 530), (1110, 655))
    arrow(draw, (810, 670), (720, 670))
    arrow(draw, (440, 670), (350, 670))
    image.save(ASSETS / "01_architettura.png")


def create_analysis_flow_diagram():
    image = Image.new("RGB", (1600, 820), "#07131B")
    draw = ImageDraw.Draw(image)
    draw.text((60, 35), "Ciclo completo di un'analisi", fill="white", font=font(42, True))
    items = [
        ("1", "Domanda + file", "CSV / Excel"),
        ("2", "Ingestion", "lettura in memoria"),
        ("3", "Validazione", "schema e qualità"),
        ("4", "Planning", "intento e colonne"),
        ("5", "Calcolo", "motore deterministico"),
        ("6", "Report", "risposta e perimetro"),
        ("7", "Intelligence", "KG e decisioni"),
        ("8", "Feedback", "voto ed esito"),
    ]
    for idx, (number, title, subtitle) in enumerate(items):
        row = idx // 4
        col = idx % 4
        x1 = 55 + col * 390
        y1 = 155 + row * 310
        xy = (x1, y1, x1 + 320, y1 + 185)
        box(draw, xy, f"{number}. {title}", subtitle)
        if col < 3:
            arrow(draw, (x1 + 320, y1 + 92), (x1 + 375, y1 + 92))
    arrow(draw, (1500, 340), (1500, 465))
    image.save(ASSETS / "02_ciclo_analisi.png")


def create_intelligence_diagram():
    image = Image.new("RGB", (1600, 780), "#061018")
    draw = ImageDraw.Draw(image)
    draw.text((60, 35), "Product Intelligence: dalla prova alla decisione", fill="white", font=font(42, True))
    stages = [
        ("Knowledge Graph", "lineage e provenance"),
        ("Governance", "struttura e consistenza"),
        ("Experience", "memoria analitica"),
        ("Recommendation", "azioni candidate"),
        ("Decision", "ranking o astensione"),
        ("Narrative", "formattazione opzionale"),
    ]
    for idx, (title, subtitle) in enumerate(stages):
        col = idx % 3
        row = idx // 3
        x1 = 95 + col * 500
        y1 = 165 + row * 290
        box(draw, (x1, y1, x1 + 390, y1 + 170), title, subtitle, fill="#112B3C" if title != "Narrative" else "#3A2C12", outline="#54D7EA" if title != "Narrative" else "#E0A22B")
        if col < 2:
            arrow(draw, (x1 + 390, y1 + 85), (x1 + 485, y1 + 85))
    arrow(draw, (1395, 335), (1395, 455))
    image.save(ASSETS / "03_product_intelligence.png")


def prepare_assets():
    ASSETS.mkdir(parents=True, exist_ok=True)
    create_architecture_diagram()
    create_analysis_flow_diagram()
    create_intelligence_diagram()


def add_figure(doc, path, caption, width=6.65, alt=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    inline = run.add_picture(str(path), width=Inches(width))
    inline._inline.docPr.set("descr", alt or caption)
    inline._inline.docPr.set("title", caption)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    set_font(cap.add_run(caption), 9, False, GRAY, True)


def label_para(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    set_font(p.add_run(label + ": "), 10.5, True, TEAL)
    set_font(p.add_run(text), 10.5, False, NAVY)


def feature(doc, title, what, does, how, users, output, limits=None):
    doc.add_heading(title, level=2)
    label_para(doc, "Cos'è", what)
    label_para(doc, "Cosa fa", does)
    label_para(doc, "Chi la usa", users)
    label_para(doc, "Come si usa", how)
    label_para(doc, "Cosa produce", output)
    if limits:
        label_para(doc, "Limiti e attenzioni", limits)


def product_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        shade(table.rows[0].cells[idx], "DCE6F1")
        set_font(table.rows[0].cells[idx].paragraphs[0].add_run(header), 9.5, True, NAVY)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_font(cells[idx].paragraphs[0].add_run(str(value)), 9.2, idx == 0, NAVY)
    set_table_geometry(table, widths)
    return table


def page_break(doc):
    doc.add_page_break()


def add_cover(doc):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("VERAXIS"), 14, True, TEAL)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Manuale di prodotto")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Funzionalità, architettura, amministrazione, API e operations")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("Guida commerciale e tecnica per vendita, onboarding e gestione del servizio"), 11, False, GRAY, True)
    doc.add_paragraph()
    callout(doc, "SCOPO DEL DOCUMENTO", "Spiegare in modo completo che cosa offre Veraxis, come funziona ogni capacità, quali ruoli la utilizzano e quali condizioni operative sono necessarie per commercializzarla in modo responsabile.", LIGHT_BLUE)
    callout(doc, "STATO DEL PRODOTTO", "Il documento descrive il perimetro attuale della private beta. Le capacità Kernel sperimentali non sono presentate come production boundary: la pipeline commerciale resta basata sul Coordinator.", LIGHT_GOLD)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("Edizione Private Beta - Luglio 2026"), 10, True, GOLD)
    page_break(doc)


def add_contents(doc):
    doc.add_heading("Indice funzionale", level=1)
    sections = [
        "1. Posizionamento del prodotto",
        "2. Utenti, ruoli e casi d'uso",
        "3. Architettura e componenti",
        "4. Accesso, organizzazioni e sicurezza",
        "5. Portale e gestione delle analisi",
        "6. Ingestion e preparazione dei dati",
        "7. Motore analitico e intenti supportati",
        "8. Ciclo asincrono, stati e cancellazione",
        "9. Pagina risultato e report",
        "10. Feedback e misurazione della qualità",
        "11. Knowledge Intelligence Workspace",
        "12. Product Intelligence e narrativa opzionale",
        "13. REST API e integrazioni",
        "14. Amministrazione, osservabilità e operations",
        "15. Lifecycle dei dati, backup e retention",
        "16. Deployment e responsabilità del cliente",
        "17. Readiness commerciale, limiti e roadmap",
        "18. Glossario e matrice delle funzionalità",
    ]
    for section in sections:
        bullet(doc, section)
    callout(doc, "COME LEGGERLO", "I capitoli 1-3 servono a vendita e decision maker; 4-12 a utenti e product owner; 13-16 a IT e amministratori; 17-18 a governance, procurement e supporto.", LIGHT_GREEN)


def build_content(doc):
    add_contents(doc)

    doc.add_heading("1. Posizionamento del prodotto", level=1)
    doc.add_paragraph("Veraxis è una piattaforma analitica multi-tenant che trasforma file CSV o Excel e domande in linguaggio naturale in analisi riproducibili, report verificabili e memoria organizzativa. Il valore commerciale nasce dall'unione di calcolo deterministico, spiegabilità, isolamento dei clienti e Knowledge Intelligence.")
    feature(doc, "Proposta di valore", "Un ambiente unico per chiedere, calcolare, verificare e riutilizzare analisi.", "Riduce il passaggio manuale tra fogli, script, report e memoria delle decisioni.", "Il cliente accede al portale, carica dati autorizzati, formula una domanda e valuta il risultato.", "Team business, operations, data analyst, amministratori e sistemi integrati.", "Report, risultati deterministici, profilo dati, intelligence, feedback e lineage.", "Non sostituisce la validazione umana né autorizza decisioni automatiche ad alto impatto.")
    doc.add_heading("Principi di prodotto", level=2)
    for item in [
        "Deterministic-first: i numeri autorevoli vengono da calcoli ripetibili.",
        "Offline-first: le funzioni core non richiedono un LLM.",
        "Tenant isolation: ogni organizzazione opera nel proprio perimetro.",
        "Evidence-first: raccomandazioni e decisioni devono avere provenance.",
        "Human feedback: la qualità viene verificata da utenti reali.",
        "Fail clearly: richieste non supportate devono produrre astensione o errore comprensibile.",
    ]:
        bullet(doc, item)

    doc.add_heading("2. Utenti, ruoli e casi d'uso", level=1)
    role_rows = [
        ("Amministratore", "Gestisce utenti e governance del tenant; può usare API amministrative e cancellare analisi."),
        ("Analista", "Carica file, avvia e annulla analisi, consulta risultati e Knowledge Intelligence."),
        ("Lettore", "Consulta risorse autorizzate senza avviare nuove elaborazioni."),
        ("Client API", "Integra Veraxis in un processo applicativo mediante token Bearer."),
        ("Operatore IT", "Gestisce deployment, health, metriche, backup, retention e incidenti."),
    ]
    product_table(doc, ("Ruolo", "Responsabilità principale"), role_rows, [2300, 7060])
    doc.add_heading("Casi d'uso commerciali", level=2)
    for item in [
        "Conteggi e distribuzioni per stato, categoria, canale o segmento.",
        "Top N, somme, medie, minimi e massimi per gruppo.",
        "Trend temporali, controllo di valori mancanti e duplicati.",
        "Individuazione di anomalie e supporto all'analisi delle cause.",
        "Creazione di un archivio consultabile di analisi, evidenze e report.",
        "Integrazione via API in portali, workflow o applicazioni del cliente.",
    ]:
        bullet(doc, item)
    callout(doc, "SEGMENTI IDEALI", "Team che lavorano con dati tabellari e desiderano risultati tracciabili senza costruire ogni volta una pipeline analitica dedicata.", LIGHT_BLUE)

    page_break(doc)
    doc.add_heading("3. Architettura e componenti", level=1)
    add_figure(doc, ASSETS / "01_architettura.png", "Figura 1 - Architettura logica della piattaforma commerciale.")
    components = [
        ("Gateway Nginx", "Espone un ingresso unico, applica limiti e header di sicurezza."),
        ("Platform API / Portal", "Autenticazione, RBAC, sessioni, API, storico e risultati."),
        ("PostgreSQL / SQLite", "Persistenza di tenant, utenti, analisi e feedback."),
        ("Redis + RQ", "Coda durevole, retry, worker separato e cancellazione cooperativa."),
        ("Coordinator", "Production boundary che orchestra ingestion, validazione, analisi e report."),
        ("Product Intelligence", "Governance, Knowledge Graph, Experience, Recommendation e Decision."),
        ("Tenant storage", "Knowledge Graph, Experience e history isolati per organizzazione."),
    ]
    for name, description in components:
        label_para(doc, name, description)
    callout(doc, "BOUNDARY DI PRODUZIONE", "Il Coordinator è il motore orchestratore in produzione. Il Kernel sperimentale viene usato per test di parità e non sostituisce ancora questo confine.", LIGHT_GOLD)

    doc.add_heading("4. Accesso, organizzazioni e sicurezza", level=1)
    add_figure(doc, USER_ASSETS / "01_accesso_registrazione.png", "Figura 2 - Accesso e registrazione dell'organizzazione.")
    feature(doc, "Registrazione organizzazione", "Creazione di un tenant e del primo amministratore.", "Genera lo spazio isolato, l'utente admin e le credenziali di sessione.", "Compilare organizzazione, email e password; in produzione la self-registration può essere disabilitata.", "Amministratore iniziale o onboarding team.", "Tenant ID, utente admin e sessione autenticata.", "In produzione è raccomandato provisioning controllato e self-registration disattivata.")
    feature(doc, "Login e logout", "Accesso tenant-aware tramite ID tenant, email e password.", "Verifica le credenziali e crea una sessione firmata HttpOnly/SameSite.", "Inserire i tre dati nel portale; usare Esci al termine della sessione.", "Tutti gli utenti del portale.", "Sessione autenticata associata a ruolo e tenant.", "HTTPS obbligatorio in produzione; password e tenant ID non vanno condivisi.")
    feature(doc, "RBAC", "Controllo accessi basato su admin, analyst e viewer.", "Limita creazione utenti, invio analisi, cancellazione e consultazione.", "Assegnare il ruolo minimo necessario durante la creazione dell'utente via API.", "Amministratori e security owner.", "Autorizzazioni coerenti applicate a ogni endpoint.", "RBAC non sostituisce processi di offboarding e revisione periodica degli accessi.")

    page_break(doc)
    doc.add_heading("5. Portale e gestione delle analisi", level=1)
    add_figure(doc, USER_ASSETS / "02_home_nuova_analisi.png", "Figura 3 - Home autenticata: nuova analisi, tenant, Knowledge Intelligence e storico.")
    feature(doc, "Nuova analisi", "Modulo principale per combinare domanda e dataset.", "Trasforma il file in record in memoria e crea un job asincrono.", "Scrivere la domanda, scegliere CSV/Excel, decidere sulla narrativa e premere Avvia analisi.", "Admin e analyst.", "ID job, stato iniziale e voce nello storico.", "Una domanda troppo generica può produrre astensione o una risposta non pertinente.")
    feature(doc, "Storico tenant", "Elenco delle analisi del tenant ordinate dalla più recente.", "Mostra data, domanda, fonte, stato, avanzamento, ID e azioni.", "Usare Apri risultato, Annulla o tornare in seguito con lo stesso account.", "Utenti autenticati secondo ruolo.", "Vista operativa e accesso persistente ai risultati.", "Non mostra dati di tenant diversi; il numero di elementi è bounded.")
    feature(doc, "Narrativa opzionale", "Flag per richiedere una riscrittura più naturale del contenuto deterministico.", "Attiva la facade narrativa soltanto quando configurata e consentita.", "Selezionare la checkbox prima dell'invio.", "Admin e analyst, secondo policy cliente.", "Testo riformattato con provenance e fallback.", "Non può introdurre nuovi numeri o fatti e non sostituisce il risultato deterministico.")

    doc.add_heading("6. Ingestion e preparazione dei dati", level=1)
    feature(doc, "Upload CSV / Excel", "Ingresso tabellare del portale.", "Legge file CSV, XLSX o XLS e costruisce un DataFrame in memoria.", "Caricare una tabella con intestazione unica, una riga per record e tipi coerenti.", "Admin e analyst.", "Dataset in memoria e metadati di fonte.", "Righe decorative, celle unite, header multipli e numeri memorizzati come testo riducono la qualità.")
    doc.add_heading("Regole di qualità prima dell'upload", level=2)
    for item in [
        "Intestazioni univoche e descrittive.",
        "Nessun subtotale o riga di titolo dentro la tabella.",
        "Date, numeri e categorie con formato coerente.",
        "Duplicati e valori mancanti compresi e documentati.",
        "Assenza di dati personali non necessari.",
        "Dimensione entro API_MAX_RECORDS e limiti del gateway.",
    ]:
        bullet(doc, item)
    callout(doc, "PERSISTENZA DEI DATI", "Le righe originali sono usate in memoria durante l'analisi ma non vengono incluse nella richiesta o nel risultato persistito dalla piattaforma. Knowledge Graph, history e log seguono lifecycle separati.", LIGHT_GREEN)

    page_break(doc)
    doc.add_heading("7. Motore analitico e intenti supportati", level=1)
    intents = [
        ("Conteggio", "count_occurrences", "Conta elementi per categoria; può includere dimensioni correlate e tabelle incrociate."),
        ("Top N", "top_n", "Classifica categorie per frequenza o metrica aggregata."),
        ("Aggregazione", "numeric_aggregation", "Somma, media, minimo, massimo o conteggio, anche per gruppo."),
        ("Trend", "time_trend", "Aggrega una metrica o i record lungo una colonna temporale."),
        ("Null", "null_detection", "Conta valori mancanti e percentuali per colonna."),
        ("Duplicati", "duplicate_detection", "Conta righe duplicate e fornisce un campione bounded."),
        ("Anomalie", "pipeline Coordinator", "Applica metodi statistici e deduplica lo stesso evento rilevato da più metodi."),
        ("Cause e raccomandazioni", "pipeline Coordinator", "Produce ipotesi supportate dalle evidenze disponibili e suggerimenti non vincolanti."),
    ]
    product_table(doc, ("Funzione", "Contratto", "Comportamento"), intents, [1800, 2500, 5060])
    doc.add_heading("Domande consigliate", level=2)
    for item in [
        "Conta i ticket per stato e canale.",
        "Mostra i primi 5 prodotti per somma dei ricavi.",
        "Calcola la media del tempo di risposta per team.",
        "Mostra il trend mensile degli ordini.",
        "Trova valori mancanti e righe duplicate.",
    ]:
        bullet(doc, item)
    callout(doc, "ASTENSIONE", "Se la richiesta non identifica un intento supportato, il motore deve dichiararla non supportata e chiedere maggiore precisione, invece di calcolare una metrica arbitraria.", LIGHT_RED)

    doc.add_heading("8. Ciclo asincrono, stati e cancellazione", level=1)
    add_figure(doc, ASSETS / "02_ciclo_analisi.png", "Figura 4 - Dalla richiesta al feedback: ciclo end-to-end di un'analisi.")
    add_figure(doc, USER_ASSETS / "03_analisi_in_elaborazione.png", "Figura 5 - Job in elaborazione con aggiornamento automatico e annullamento.")
    state_rows = [
        ("queued", "In coda", "Richiesta persistita, in attesa del worker."),
        ("processing", "In elaborazione", "Coordinator in esecuzione; progress aggiornabile."),
        ("cancelling", "Annullamento", "Richiesta di stop registrata."),
        ("cancelled", "Annullata", "Job terminato senza pubblicare il risultato."),
        ("completed", "Completata", "Risultato verificato e persistito."),
        ("failed", "Non riuscita", "Errore registrato; il report non viene presentato come valido."),
    ]
    product_table(doc, ("Stato API", "Etichetta", "Significato"), state_rows, [1900, 2200, 5260])
    feature(doc, "Cancellazione cooperativa", "Comando di arresto per job in coda o in elaborazione.", "Imposta cancel_requested e interrompe il worker al successivo punto di controllo sicuro.", "Dal portale usare Annulla; via API usare POST /analyses/{id}/cancel.", "Admin e analyst.", "Stato cancelled o cancelling.", "Non equivale a cancellazione dei dati persistiti; quella è un'operazione amministrativa separata.")

    page_break(doc)
    doc.add_heading("9. Pagina risultato e report", level=1)
    add_figure(doc, USER_ASSETS / "05_risultato_top.png", "Figura 6 - Risultato: metadati, report, calcolo deterministico e accesso al Knowledge Graph.")
    feature(doc, "Hero e metadati", "Intestazione dell'analisi.", "Mostra domanda, ID, stato, avanzamento, fonte e data.", "Confrontare questi dati con la richiesta originale prima di leggere i numeri.", "Tutti gli utenti autorizzati.", "Contesto minimo per audit e supporto.")
    feature(doc, "Report dell'analisi", "Risposta business answer-first in Markdown.", "Riassume obiettivo, perimetro, risultati, qualità dati e raccomandazioni pertinenti.", "Leggere prima la risposta, poi verificare tabelle e caveat.", "Business user e analyst.", "Testo scaricabile in formato Markdown.", "Il report può essere formalmente leggibile ma non pertinente: va sempre confrontato con la domanda.")
    feature(doc, "Risultati deterministici", "Payload tecnico dei calcoli ripetibili.", "Espone piano, colonne, aggregazioni, valori, status e provenance.", "Usarlo per audit, debug e riconciliazione dei numeri.", "Analyst, QA, supporto e integrazioni.", "JSON visualizzato in pannello bounded.")
    feature(doc, "Perimetro dati", "Profilo del DataFrame elaborato.", "Mostra righe, colonne, tipi, null, duplicati, statistiche e correlazioni disponibili.", "Verificare che il dataset letto coincida con quello previsto.", "Analyst e data owner.", "Profilo strutturale senza dump delle righe originali.")
    feature(doc, "Intelligenza di prodotto", "Payload integrato di governance e decision support.", "Raccoglie consistenza, Knowledge Graph, Experience, raccomandazioni, decisione, narrativa e osservabilità.", "Esaminare status, evidenze, rischio e provenance.", "Analyst avanzato, product owner, audit.", "Payload JSON e sezioni nel report.", "Una failure di intelligence non invalida automaticamente un calcolo analitico valido.")
    feature(doc, "Download Markdown", "Esportazione testuale del report.", "Scarica il report dell'analisi senza il payload tecnico completo.", "Premere Scarica report Markdown.", "Utenti autorizzati.", "File .md archiviabile e versionabile.", "La condivisione segue le policy del cliente.")

    doc.add_heading("10. Feedback e misurazione della qualità", level=1)
    add_figure(doc, USER_ASSETS / "05_feedback.png", "Figura 7 - Valutazione 1-5, esito e nota facoltativa.", width=5.7)
    feature(doc, "Feedback tenant-scoped", "Valutazione associata a utente, tenant e analisi.", "Registra rating 1-5, esito correct/partial/incorrect e nota fino a 1.000 caratteri.", "Valutare il risultato dopo un controllo indipendente e premere Salva feedback.", "Utenti autenticati che conoscono il dataset.", "Record aggiornabile dallo stesso utente e metriche aggregate.", "Non premiare la grafica; valutare pertinenza, colonne, aggregazioni e numeri.")
    doc.add_heading("Metriche aggregate", level=2)
    for item in [
        "numero totale di feedback",
        "rating medio",
        "conteggi per esito",
        "analisi persistite per stato",
        "counter di submission, completamenti e failure del processo corrente",
    ]:
        bullet(doc, item)
    callout(doc, "PRIVACY DELLE METRICHE", "L'endpoint /metrics espone aggregati senza email, note, prompt o identificativi tenant nelle label.", LIGHT_GREEN)

    page_break(doc)
    doc.add_heading("11. Knowledge Intelligence Workspace", level=1)
    add_figure(doc, USER_ASSETS / "06_knowledge_graph.png", "Figura 8 - Workspace tenant-scoped della conoscenza e delle decisioni.")
    knowledge_features = [
        ("Metriche del grafo", "Nodi, relazioni, tipi, esperienze, analisi e governance."),
        ("Ricerca semantica", "Filtra label, ID e proprietà dei nodi."),
        ("Filtro per tipo", "Riduce la vista a analysis_run, dataset, colonne, insight o report."),
        ("Relazioni dominanti", "Mostra le connessioni più frequenti nel grafo."),
        ("Mappa interattiva", "Visualizza nodi e archi con zoom e selezione."),
        ("Dettaglio nodo", "Espone proprietà, provenance e vicini del nodo selezionato."),
        ("Console deterministica", "Risponde a domande supportate usando solo evidenze del grafo."),
        ("Experience", "Mostra pattern analitici consolidati nel tempo."),
        ("Raccomandazioni", "Presenta azioni candidate con score ed evidenze."),
        ("Decisioni", "Mostra ranking, selezione o astensione della policy."),
        ("Qualità", "Espone validazione strutturale, issue, quarantena e possibilità di consumo/scrittura."),
        ("Cronologia", "Collega le analisi completate alla memoria del tenant."),
        ("Esportazione", "Scarica una proiezione JSON bounded del workspace."),
    ]
    for title, description in knowledge_features:
        label_para(doc, title, description)
    callout(doc, "FUNZIONA OFFLINE", "La query deterministica e la visualizzazione non richiedono un LLM o un graph database esterno. Le risposte sono limitate alle evidenze già presenti.", LIGHT_BLUE)

    doc.add_heading("12. Product Intelligence e narrativa opzionale", level=1)
    add_figure(doc, ASSETS / "03_product_intelligence.png", "Figura 9 - Pipeline di intelligence evidence-based.", width=5.9)
    stages = [
        ("Knowledge Graph", "Mappa dataset, colonne, analisi, insight e report con provenance."),
        ("Governance", "Valida struttura, naming, identità, riferimenti e compatibilità dello schema."),
        ("Consistency", "Blocca il riuso quando le evidenze non superano i gate semantici."),
        ("Experience", "Consolida pattern analitici riutilizzabili dal tenant."),
        ("Recommendation", "Raccoglie e ordina azioni candidate bounded."),
        ("Decision", "Seleziona un'opzione supportata oppure si astiene."),
        ("Narrative", "Formatta i fatti ammessi senza sostituire il contenuto deterministico."),
        ("Observability", "Registra run ID, stage, durata, status ed error type."),
    ]
    for name, description in stages:
        label_para(doc, name, description)
    callout(doc, "FALLBACK DETERMINISTICO", "Se la narrativa è disabilitata, offline, oltre budget, non disponibile o sopra i limiti, viene restituito il testo deterministico. Il risultato dichiara sempre se un modello è stato usato.", LIGHT_GOLD)

    doc.add_heading("13. REST API e integrazioni", level=1)
    doc.add_paragraph("La REST API consente di integrare Veraxis in applicazioni e workflow del cliente. L'autenticazione applicativa usa token Bearer firmati e time-bounded.")
    api_rows = [
        ("POST", "/api/v1/auth/register", "Provisioning tenant, se abilitato"),
        ("POST", "/api/v1/auth/login", "Login con tenant ID, email e password"),
        ("POST", "/api/v1/users", "Creazione utente; solo admin"),
        ("POST", "/api/v1/analyses", "Invio record e domanda; admin/analyst"),
        ("GET", "/api/v1/analyses", "Lista analisi del tenant"),
        ("GET", "/api/v1/analyses/{id}", "Dettaglio tenant-scoped"),
        ("POST", "/api/v1/analyses/{id}/cancel", "Cancellazione cooperativa"),
        ("POST", "/api/v1/analyses/{id}/feedback", "Rating, esito e nota"),
        ("DELETE", "/api/v1/analyses/{id}", "Eliminazione amministrativa"),
        ("GET", "/api/v1/knowledge", "Read model Knowledge Intelligence"),
        ("POST", "/api/v1/knowledge/query", "Query deterministica del grafo"),
        ("GET", "/api/v1/openapi.json", "Contratto OpenAPI"),
    ]
    product_table(doc, ("Metodo", "Endpoint", "Funzione"), api_rows, [1200, 4300, 3860])
    doc.add_heading("Contratto di invio analisi", level=2)
    label_para(doc, "Input minimo", "description non vuota e records come lista non vuota di oggetti JSON.")
    label_para(doc, "Risposta", "HTTP 202 con ID, stato queued e backend della coda.")
    label_para(doc, "Polling", "GET sul dettaglio finché lo stato è completed, failed o cancelled.")
    label_para(doc, "Isolamento", "Un token non può leggere o modificare analisi di un altro tenant.")
    label_para(doc, "Compatibilità sorgente", "I record JSON vengono normalizzati verso un adapter eseguibile; le label client restano provenance, non nomi di adapter arbitrari.")

    doc.add_heading("14. Amministrazione, osservabilità e operations", level=1)
    feature(doc, "Health live", "Endpoint /health/live per verificare che il processo risponda.", "Restituisce lo stato di vita dell'applicazione.", "Configurare probe infrastrutturali periodici.", "Orchestratore, load balancer, SRE.", "HTTP 200 quando il processo è vivo.")
    feature(doc, "Health ready", "Endpoint /health/ready per verificare dipendenze e capacità di servire traffico.", "Controlla la readiness della piattaforma.", "Usarlo in deployment e bilanciamento.", "Kubernetes/Compose, SRE, supporto.", "HTTP 200 ready o segnale di indisponibilità.")
    feature(doc, "Prometheus metrics", "Endpoint testuale /metrics.", "Espone counter di analisi, feedback aggregati e stati persistiti.", "Scrapare con Prometheus o un collector compatibile.", "Operations e product analytics.", "Serie aggregate senza dati grezzi.", "Proteggere l'endpoint a livello di rete o gateway in produzione.")
    feature(doc, "Log strutturati", "Log rotanti dell'applicazione e telemetria per stage.", "Registra run ID, durata, status ed error type senza righe originali.", "Centralizzare, applicare retention e alert su failure/latency.", "SRE, supporto e incident manager.", "Audit operativo e diagnosi.")
    feature(doc, "Utenti", "Provisioning degli account tramite API amministrativa.", "Crea utenti nel tenant con ruolo esplicito e password hashata.", "Chiamare POST /api/v1/users con token admin.", "Amministratore del tenant.", "ID utente e associazione tenant.", "L'attuale portale non espone una schermata completa di user management.")

    doc.add_heading("15. Lifecycle dei dati, backup e retention", level=1)
    feature(doc, "Cancellazione analisi", "DELETE amministrativa tenant-scoped.", "Elimina analisi e feedback associati.", "Usare DELETE /api/v1/analyses/{id} con ruolo admin.", "Admin o processo di privacy operations.", "HTTP 204 quando completata.", "Knowledge Graph, Experience, log e backup hanno lifecycle separati.")
    feature(doc, "Retention", "Policy batch per job terminali più vecchi della soglia.", "Identifica candidati in dry-run e li elimina solo con --apply.", "Eseguire scripts/enforce_retention.py --days N, verificare, poi aggiungere --apply.", "Operations e data governance.", "Conteggio candidati/eliminati.", "Non applicare senza policy approvata e backup coerente.")
    feature(doc, "Backup SQLite", "Copia consistente del database locale.", "Produce un file di backup fuori dal database attivo.", "Usare scripts/backup_platform.py --output-dir ... .", "Operations.", "Database di backup versionabile esternamente.")
    feature(doc, "Restore SQLite", "Ripristino esplicito da backup.", "Sostituisce il database soltanto con conferma.", "Provare in ambiente isolato con scripts/restore_platform.py ... --confirm.", "Operations e disaster recovery.", "Database ripristinato e verificabile.", "PostgreSQL richiede pg_dump/pg_restore e procedure del cliente.")
    callout(doc, "PROVA DI RESTORE", "Un backup non è considerato sufficiente finché una prova isolata non verifica schema, tenant, analisi, feedback e readiness.", LIGHT_RED)

    doc.add_heading("16. Deployment e responsabilità del cliente", level=1)
    deploy_rows = [
        ("Nginx", "Gateway, limiti, security header, TLS termination o upstream TLS."),
        ("Gunicorn API", "Servizio web multi-worker non-root."),
        ("RQ worker", "Elaborazione asincrona separata dall'API."),
        ("Redis", "Coda con persistenza AOF e no-eviction."),
        ("PostgreSQL", "Persistenza multi-user e migrazioni serializzate."),
        ("Volumi", "Database, tenant storage, log e backup secondo policy."),
        ("Secrets", "DATABASE_URL e PLATFORM_AUTH_SECRET fuori dal repository."),
    ]
    product_table(doc, ("Componente", "Responsabilità"), deploy_rows, [2400, 6960])
    doc.add_heading("Responsabilità operative del cliente", level=2)
    for item in [
        "Fornire HTTPS, DNS, firewall e accesso di rete controllato.",
        "Gestire segreti, rotazione credenziali e offboarding utenti.",
        "Definire categorie di dati consentite e policy di retention.",
        "Monitorare health, metriche, coda, database, storage e backup.",
        "Stabilire canale di supporto, escalation e gestione incidenti.",
        "Validare requisiti legali e privacy del proprio settore.",
    ]:
        bullet(doc, item)

    page_break(doc)
    doc.add_heading("17. Readiness commerciale, limiti e roadmap", level=1)
    doc.add_paragraph("Il framework tecnico della private beta è implementato. La commercializzazione generale richiede evidenze sufficienti di qualità sul campo, oltre ai test automatici.")
    gate_rows = [
        ("Campione funzionale", "30 casi / almeno 3 domini", "Superato: 30 casi, 6 domini"),
        ("Feedback verificati", "almeno 10", "Da raccogliere con utenti reali"),
        ("Accuratezza", "almeno 80% correct", "Calcolata dopo i feedback"),
        ("Concorrenza", "almeno 5, error rate <=2%", "Gate tecnico superato"),
        ("Isolamento / operations", "tutti superati", "Gate tecnici superati"),
        ("Bug critici", "zero", "Richiesto prima dell'apertura"),
    ]
    product_table(doc, ("Gate", "Soglia", "Stato"), gate_rows, [2600, 3300, 3460])
    doc.add_heading("Limiti dichiarati", level=2)
    for item in [
        "Il portale beta privilegia CSV/Excel; connettori enterprise non sono parte del flusso commerciale attuale.",
        "Il user management completo è API-first.",
        "La cancellazione dei job è cooperativa, non un kill immediato del processo.",
        "Knowledge Graph ed Experience richiedono policy di lifecycle separate dal record analisi.",
        "La narrativa opzionale è una funzione di formattazione e può non essere disponibile.",
        "Il Kernel resta sperimentale fino alla parità documentata con il Coordinator.",
        "Decisioni ad alto impatto richiedono sempre revisione umana e governance specifica.",
    ]:
        bullet(doc, item)
    callout(doc, "CRITERIO DI VENDITA", "Presentare Veraxis come piattaforma analitica verificabile e assistita, non come sostituto autonomo del data analyst o del decision maker.", LIGHT_GOLD)

    doc.add_heading("18. Glossario e matrice delle funzionalità", level=1)
    glossary = [
        ("Coordinator", "Orchestratore della pipeline di produzione."),
        ("Kernel", "Runtime sperimentale modulare, non ancora production boundary."),
        ("Deterministico", "Calcolo ripetibile con stesso input e stesse regole."),
        ("Tenant", "Perimetro isolato di una singola organizzazione."),
        ("RBAC", "Autorizzazioni assegnate in base al ruolo."),
        ("Knowledge Graph", "Rete di analisi, dataset, colonne, insight e report."),
        ("Experience", "Memoria di pattern analitici riutilizzabili."),
        ("Recommendation", "Azione candidata ordinata per evidenze e policy."),
        ("Decision", "Selezione o astensione registrata con provenance."),
        ("Narrative", "Formattazione opzionale dei fatti deterministici."),
        ("Provenance", "Traccia dell'origine e delle trasformazioni di un'evidenza."),
        ("Readiness", "Insieme di gate che autorizza un rilascio o una beta."),
    ]
    for term, definition in glossary:
        label_para(doc, term, definition)
    doc.add_heading("Matrice rapida", level=2)
    matrix = [
        ("Portale", "Utente", "Operativa", "Upload, storico, risultati"),
        ("API", "Sistemi / admin", "Operativa", "Integrazione e amministrazione"),
        ("Motore deterministico", "Analyst", "Operativa", "Calcoli autorevoli"),
        ("Knowledge Intelligence", "Analyst / decision maker", "Operativa", "Memoria e provenance"),
        ("Narrativa", "Business user", "Opzionale", "Formattazione"),
        ("Kernel", "Engineering", "Sperimentale", "Parità e capability future"),
    ]
    product_table(doc, ("Area", "Utente", "Stato", "Scopo"), matrix, [1900, 2200, 1800, 3460])
    callout(doc, "DOCUMENTO DI CONSEGNA", "Questo manuale può accompagnare demo, onboarding cliente, security review e handoff operativo. Configurazioni, soglie e responsabilità devono essere adattate al contratto e all'ambiente del cliente.", LIGHT_GREEN)


def main():
    prepare_assets()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure(doc)
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header.add_run("VERAXIS  |  MANUALE DI PRODOTTO"), 8.5, True, GRAY)
    footer = section.footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("Manuale di prodotto - Private Beta  |  Pagina "), 8.5, False, GRAY)
    add_page_field(footer)
    add_cover(doc)
    build_content(doc)
    core = doc.core_properties
    core.title = "Manuale di prodotto Veraxis"
    core.subject = "Funzionalità, architettura, amministrazione, API e operations"
    core.author = "Veraxis"
    core.keywords = "Veraxis, prodotto, analytics, Knowledge Graph, API, operations"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
