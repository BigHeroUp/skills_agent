"""Build the illustrated Veraxis private-beta user manual."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ASSETS = ROOT / "docs" / "manual_assets"
OUTPUT = ROOT / "docs" / "deliverables" / "Manuale_utente_Veraxis_Beta.docx"
NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
TEAL = RGBColor(0, 124, 137)
GOLD = RGBColor(180, 126, 20)
GRAY = RGBColor(85, 96, 110)
LIGHT_BLUE = "EAF2F8"
LIGHT_GOLD = "FFF4D6"
LIGHT_RED = "FCE8E6"
LIGHT_GREEN = "E6F4EA"
WHITE = RGBColor(255, 255, 255)


def set_font(run, size=11, bold=False, color=None, italic=False):
    run.font.name = "Aptos"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Aptos")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Aptos")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if len(table.rows) > 1:
        tr_pr = table.rows[0]._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:tblHeader")) is None:
            header = OxmlElement("w:tblHeader")
            header.set(qn("w:val"), "true")
            tr_pr.append(header)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def configure(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = NAVY
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    tokens = {
        "Title": (30, NAVY, 0, 8),
        "Subtitle": (15, GRAY, 0, 18),
        "Heading 1": (18, BLUE, 18, 8),
        "Heading 2": (14, NAVY, 14, 6),
        "Heading 3": (11.5, TEAL, 10, 4),
    }
    for name, (size, color, before, after) in tokens.items():
        style = doc.styles[name]
        style.font.name = "Aptos Display" if name != "Normal" else "Aptos"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Aptos"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header.add_run("VERAXIS  |  GUIDA PRIVATE BETA"), 8.5, True, GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("Manuale utente - versione beta  |  Pagina "), 8.5, False, GRAY)
    add_page_field(footer)


def add_cover(doc):
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("VERAXIS"), 13, True, TEAL)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Manuale utente illustrato")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Guida completa per partecipanti alla private beta senza esperienza precedente")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("Dall'accesso al feedback: ogni passaggio spiegato in modo semplice"), 11, False, GRAY, True)
    doc.add_paragraph()
    callout(doc, "A CHI SERVE", "A persone che non hanno mai usato Veraxis e devono caricare un file, formulare una domanda, leggere il risultato e valutarne l'accuratezza.", LIGHT_BLUE)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("Private beta - Luglio 2026"), 10, True, GOLD)
    doc.add_page_break()


def callout(doc, label, text, fill=LIGHT_BLUE):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "160")
    ind.set(qn("w:right"), "160")
    p_pr.append(ind)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    set_font(p.add_run(label + "\n"), 9, True, TEAL if fill != LIGHT_RED else RGBColor(155, 28, 28))
    set_font(p.add_run(text), 10.5, False, NAVY)


def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix and text.startswith(bold_prefix):
        set_font(p.add_run(bold_prefix), 10.5, True, NAVY)
        set_font(p.add_run(text[len(bold_prefix):]), 10.5, False, NAVY)
    else:
        set_font(p.add_run(text), 10.5, False, NAVY)
    return p


def step(doc, number, title, body):
    p = doc.add_paragraph(style="Heading 2")
    set_font(p.add_run(f"Passo {number} - {title}"), 14, True, NAVY)
    p = doc.add_paragraph(body)
    p.paragraph_format.space_after = Pt(8)


def add_figure(doc, filename, caption, width=6.65, alt=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    inline = run.add_picture(str(ASSETS / filename), width=Inches(width))
    doc_pr = inline._inline.docPr
    doc_pr.set("descr", alt or caption)
    doc_pr.set("title", caption)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    set_font(cap.add_run(caption), 9, False, GRAY, True)


def add_contents(doc):
    doc.add_heading("Come usare questo manuale", level=1)
    doc.add_paragraph("Segui i capitoli nell'ordine indicato durante la prima prova. In seguito puoi usare le sezioni finali come riferimento rapido.")
    entries = [
        "1. Che cos'è Veraxis e cosa farà il tester",
        "2. Prima di iniziare: credenziali, browser e file",
        "3. Registrazione e accesso",
        "4. Preparazione corretta del file",
        "5. Avvio della prima analisi",
        "6. Scrivere una buona domanda",
        "7. Attesa, avanzamento e annullamento",
        "8. Lettura del risultato",
        "9. Valutazione e feedback beta",
        "10. Download, storico e Knowledge Graph",
        "11. Sicurezza e comportamento responsabile",
        "12. Risoluzione dei problemi",
        "13. Checklist della sessione beta e glossario",
    ]
    for entry in entries:
        bullet(doc, entry)
    callout(doc, "REGOLA D'ORO", "Non serve conoscere statistica o programmazione. Devi però conoscere il significato del tuo file e controllare se la risposta soddisfa davvero la domanda che hai scritto.", LIGHT_GOLD)


def add_manual(doc):
    add_contents(doc)

    doc.add_heading("1. Che cos'è Veraxis e cosa farà il tester", level=1)
    doc.add_paragraph("Veraxis è una piattaforma che analizza dati contenuti in file CSV o Excel. L'utente carica un file, scrive una domanda in linguaggio naturale e riceve un report con risultati, controlli sul perimetro dei dati e informazioni tecniche verificabili.")
    callout(doc, "NOTA SUL NOME", "Nell'interfaccia della beta potresti vedere la dicitura “Skills Agent”. È il nome tecnico attualmente mostrato dal portale; in questo manuale lo chiamiamo Veraxis.")
    doc.add_heading("Il tuo compito nella private beta", level=2)
    for text in [
        "Usare un file che conosci e che sei autorizzato a caricare.",
        "Scrivere una domanda precisa e verificabile.",
        "Attendere la conclusione dell'analisi senza chiudere forzatamente il processo.",
        "Confrontare il report con il file originale o con un calcolo di controllo.",
        "Salvare un feedback onesto: Corretto, Parziale oppure Errato.",
    ]:
        bullet(doc, text)
    callout(doc, "COSA NON FA VERAXIS", "Non sostituisce il giudizio professionale, non corregge automaticamente dati sbagliati e non rende vera una conclusione solo perché è presentata in un report ben formattato.", LIGHT_RED)

    doc.add_heading("2. Prima di iniziare", level=1)
    doc.add_paragraph("Chiedi al responsabile della beta l'indirizzo web del portale e le modalità di accesso. Tieni a disposizione un browser aggiornato e un file di prova piccolo, comprensibile e privo di dati non autorizzati.")
    doc.add_heading("Controllo preliminare", level=2)
    for text in [
        "Hai ricevuto il link corretto della piattaforma.",
        "Conosci la tua email, la password e, se richiesto, l'ID tenant.",
        "Il file è CSV, XLSX o XLS e non supera i limiti comunicati dal responsabile.",
        "Sai cosa rappresenta ogni riga e quali colonne servono per rispondere alla domanda.",
        "Possiedi un risultato atteso o un modo semplice per verificare la risposta.",
    ]:
        bullet(doc, text)
    callout(doc, "TENANT IN PAROLE SEMPLICI", "Il tenant è lo spazio riservato alla tua organizzazione. Analisi, utenti e memoria restano separati da quelli delle altre organizzazioni. Conserva l'ID tenant: può servire per accedere.")

    doc.add_heading("3. Registrazione e accesso", level=1)
    add_figure(doc, "01_accesso_registrazione.png", "Figura 1 - Pagina iniziale: accesso a sinistra e registrazione dell'organizzazione a destra.")
    step(doc, 1, "Apri il portale", "Apri il link ricevuto. Verifica che la pagina mostri i moduli “Accedi” e, quando abilitato, “Registra un'organizzazione”.")
    step(doc, 2, "Registrati solo se autorizzato", "Nel riquadro di registrazione inserisci il nome dell'organizzazione, l'email dell'amministratore e una password robusta. Premi “Crea utenza”. Se il responsabile ha già creato l'account, non registrare una seconda organizzazione.")
    step(doc, 3, "Conserva l'ID tenant", "Dopo la registrazione compare un codice lungo. Copialo in un luogo sicuro: identifica lo spazio della tua organizzazione.")
    step(doc, 4, "Accedi nelle sessioni successive", "Inserisci ID tenant, email e password nel riquadro “Accedi”, quindi premi il pulsante. Non condividere password o tenant ID in chat, ticket pubblici o screenshot.")

    doc.add_heading("4. Preparazione corretta del file", level=1)
    doc.add_paragraph("La qualità dell'analisi dipende dalla struttura del file. Prima del caricamento aprilo e controlla questi aspetti.")
    for text in [
        "Una sola riga di intestazione, con nomi chiari come data, regione, ricavi, stato.",
        "Una riga per ogni entità o evento; evita subtotali, titoli decorativi e righe vuote prima dell'intestazione.",
        "Date scritte in modo coerente; numeri memorizzati come numeri, non come testo con simboli misti.",
        "Categorie coerenti: per esempio non mescolare Attivo, attivo e ATTIVO se significano la stessa cosa.",
        "Nessuna formula con errori visibili e nessuna cella unita nella tabella dati.",
        "Solo dati necessari alla prova; rimuovi informazioni personali o sensibili non indispensabili.",
    ]:
        bullet(doc, text)
    callout(doc, "FILE DI ESEMPIO", "Insieme al manuale è disponibile esempio_vendite.csv. È sintetico e serve soltanto per imparare il flusso; non contiene dati reali.", LIGHT_GREEN)

    doc.add_heading("5. Avvio della prima analisi", level=1)
    add_figure(doc, "02_home_nuova_analisi.png", "Figura 2 - Home del portale: domanda, selezione del file, narrativa opzionale e storico.")
    step(doc, 1, "Scrivi la domanda", "Nel campo “Domanda di business” descrivi una sola richiesta principale. Inizia con un verbo operativo: Conta, Calcola, Confronta, Mostra, Trova oppure Verifica.")
    step(doc, 2, "Scegli il file", "Premi “Choose File” o “Scegli file”, seleziona il CSV o Excel e controlla che il nome del file compaia nel browser.")
    step(doc, 3, "Decidi sulla narrativa opzionale", "Lascia disattivata l'opzione nella prima prova, salvo indicazioni del responsabile. La parte deterministica rimane la fonte autorevole dei calcoli; la narrativa serve a rendere il testo più leggibile quando il servizio è configurato.")
    step(doc, 4, "Avvia", "Premi “Avvia analisi” una sola volta. Un messaggio conferma l'invio e assegna un ID all'analisi.")

    doc.add_heading("6. Scrivere una buona domanda", level=1)
    doc.add_paragraph("Una domanda efficace contiene l'operazione, la misura, il raggruppamento e, quando serve, il periodo o il filtro.")
    examples = [
        ("Buona", "Conta i contratti per stato e mostra lo stato dell'antenna per ciascun gruppo."),
        ("Buona", "Calcola la somma dei ricavi per regione nel primo trimestre."),
        ("Buona", "Trova le righe duplicate e indica quante sono."),
        ("Debole", "Analizza tutto."),
        ("Debole", "Dimmi qualcosa di interessante."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [1500, 7860])
    for i, text in enumerate(("Tipo", "Esempio")):
        shade(table.rows[0].cells[i], "DCE6F1")
        set_font(table.rows[0].cells[i].paragraphs[0].add_run(text), 10, True, NAVY)
    for kind, example in examples:
        cells = table.add_row().cells
        shade(cells[0], LIGHT_GREEN if kind == "Buona" else LIGHT_RED)
        set_font(cells[0].paragraphs[0].add_run(kind), 10, True, NAVY)
        set_font(cells[1].paragraphs[0].add_run(example), 10, False, NAVY)
    set_table_geometry(table, [1500, 7860])
    callout(doc, "UNA DOMANDA ALLA VOLTA", "Se chiedi contemporaneamente dieci analisi diverse, diventa difficile capire quale parte sia stata interpretata correttamente. Parti da una domanda semplice e aggiungi complessità nelle prove successive.", LIGHT_GOLD)

    doc.add_heading("7. Attesa, avanzamento e annullamento", level=1)
    add_figure(doc, "03_analisi_in_elaborazione.png", "Figura 3 - Lo storico mostra l'analisi in elaborazione e rende disponibile il comando Annulla.")
    doc.add_paragraph("Durante l'elaborazione la pagina si aggiorna automaticamente. Puoi lasciare aperto il portale oppure tornare più tardi con lo stesso account.")
    states = [
        ("In coda", "La richiesta è stata ricevuta e attende un processo disponibile."),
        ("In elaborazione", "Veraxis sta leggendo il file e producendo il risultato."),
        ("Completata", "Il risultato è disponibile. Aprilo e verificalo."),
        ("Non riuscita", "L'analisi non è stata completata. Leggi l'errore e consulta il troubleshooting."),
        ("Annullata", "L'utente o il sistema ha interrotto la richiesta."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2200, 7160])
    for i, text in enumerate(("Stato", "Significato")):
        shade(table.rows[0].cells[i], "DCE6F1")
        set_font(table.rows[0].cells[i].paragraphs[0].add_run(text), 10, True, NAVY)
    for state, meaning in states:
        cells = table.add_row().cells
        set_font(cells[0].paragraphs[0].add_run(state), 10, True, NAVY)
        set_font(cells[1].paragraphs[0].add_run(meaning), 10, False, NAVY)
    set_table_geometry(table, [2200, 7160])
    doc.add_paragraph("Usa “Annulla” solo se hai caricato il file sbagliato, scritto una domanda errata o non vuoi più proseguire. Non premere ripetutamente Avvia o Aggiorna mentre il job è in corso.")

    doc.add_heading("8. Lettura del risultato", level=1)
    add_figure(doc, "04_storico_completato.png", "Figura 4 - Quando lo stato è Completata, usa Apri risultato.")
    add_figure(doc, "05_risultato_top.png", "Figura 5 - Parte superiore della pagina risultato: obiettivo, stato, fonte, download e report.")
    doc.add_heading("Leggi in questo ordine", level=2)
    for text in [
        "Obiettivo: deve corrispondere alla domanda che hai scritto.",
        "Perimetro: controlla numero di righe, colonne e fonte.",
        "Risposta o Executive Summary: individua il numero o la tabella che risponde alla domanda.",
        "Qualità dati e anomalie: verifica se mancano valori o se esistono limiti importanti.",
        "Risultati deterministici: sono il dettaglio tecnico dei calcoli e aiutano a verificare colonne e aggregazioni.",
        "Raccomandazioni: trattale come suggerimenti da valutare, non come decisioni automatiche.",
    ]:
        bullet(doc, text)
    callout(doc, "CONTROLLO ESSENZIALE", "Un report può essere leggibile ma rispondere a una domanda diversa. Se chiedi un conteggio per regione e il risultato mostra soltanto una media generale, l'esito non è Corretto.", LIGHT_RED)

    doc.add_heading("9. Valutazione e feedback beta", level=1)
    add_figure(doc, "05_feedback.png", "Figura 6 - Modulo di feedback al fondo della pagina risultato.", width=5.7)
    doc.add_paragraph("Il feedback è la parte più importante della private beta. Salvalo soltanto dopo aver confrontato il risultato con il file o con un calcolo indipendente.")
    doc.add_heading("Come scegliere l'esito", level=2)
    verdicts = [
        ("Corretto", "Risponde alla domanda, usa le colonne giuste e i valori principali coincidono con il controllo."),
        ("Parziale", "Una parte è utile e corretta, ma manca un raggruppamento, un filtro, una misura o una spiegazione richiesta."),
        ("Errato", "Risponde a un'altra domanda, usa colonne o aggregazioni sbagliate, presenta numeri non verificabili oppure perde il dataset."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [1900, 7460])
    for i, text in enumerate(("Esito", "Quando usarlo")):
        shade(table.rows[0].cells[i], "DCE6F1")
        set_font(table.rows[0].cells[i].paragraphs[0].add_run(text), 10, True, NAVY)
    for verdict, meaning in verdicts:
        cells = table.add_row().cells
        shade(cells[0], LIGHT_GREEN if verdict == "Corretto" else LIGHT_GOLD if verdict == "Parziale" else LIGHT_RED)
        set_font(cells[0].paragraphs[0].add_run(verdict), 10, True, NAVY)
        set_font(cells[1].paragraphs[0].add_run(meaning), 10, False, NAVY)
    set_table_geometry(table, [1900, 7460])
    doc.add_heading("Come assegnare il voto", level=2)
    for text in [
        "5 - Eccellente: corretto, completo, chiaro e immediatamente utilizzabile.",
        "4 - Buono: corretto con piccole imperfezioni di forma o dettaglio.",
        "3 - Sufficiente: parzialmente utile, richiede integrazioni o controlli importanti.",
        "2 - Debole: contiene pochi elementi utili e problemi sostanziali.",
        "1 - Errato: non risponde alla domanda o presenta risultati materialmente sbagliati.",
    ]:
        bullet(doc, text)
    doc.add_heading("Scrivere una nota utile", level=2)
    doc.add_paragraph("Una buona nota indica cosa è stato richiesto, cosa è successo e quale sarebbe stato il risultato corretto. Esempio: “Ho chiesto il conteggio per regione; il report ha calcolato la media dei ricavi senza raggruppamento. Atteso: Nord 3, Sud 3”.")
    callout(doc, "NON PREMIARE LA GRAFICA", "Valuta accuratezza e pertinenza. Un report elegante ma numericamente sbagliato deve essere segnato Parziale o Errato.", LIGHT_RED)

    doc.add_heading("10. Download, storico e Knowledge Graph", level=1)
    doc.add_heading("Scaricare il report", level=2)
    doc.add_paragraph("Nella pagina risultato premi “Scarica report Markdown”. Il file .md è un documento di testo strutturato: può essere aperto con un editor, archiviato come evidenza o condiviso secondo le regole della tua organizzazione.")
    doc.add_heading("Riaprire un'analisi", level=2)
    doc.add_paragraph("Torna al portale e usa la tabella Storico analisi. Ogni riga mostra data, domanda, fonte, stato, avanzamento, ID e azioni disponibili.")
    add_figure(doc, "06_knowledge_graph.png", "Figura 7 - Spazio della conoscenza: metriche, filtri, mappa, interrogazione e cronologia.")
    doc.add_heading("Usare lo spazio della conoscenza", level=2)
    for text in [
        "I riquadri superiori riassumono nodi, relazioni, tipi, esperienze e analisi.",
        "Ricerca semantica e Tipo nodo restringono gli elementi mostrati.",
        "La mappa collega analisi, dataset, colonne, insight e report.",
        "La console risponde soltanto con evidenze presenti nella memoria del tenant.",
        "Esporta scarica una rappresentazione tecnica della conoscenza: usala solo se richiesta dal responsabile.",
    ]:
        bullet(doc, text)

    doc.add_heading("11. Sicurezza e comportamento responsabile", level=1)
    for text in [
        "Carica soltanto dati che sei autorizzato a utilizzare.",
        "Per la beta preferisci dataset sintetici, anonimizzati o a basso rischio.",
        "Non inserire password, token, segreti aziendali o dati personali nella domanda.",
        "Non condividere link, tenant ID, screenshot o report fuori dal gruppo autorizzato.",
        "Esci dal portale quando usi un computer condiviso.",
        "Segnala immediatamente accessi anomali, dati di altri tenant o risultati che sembrano provenire da file diversi.",
    ]:
        bullet(doc, text)
    callout(doc, "STOP IMMEDIATO", "Se visualizzi dati appartenenti a un'altra organizzazione, interrompi la prova, non scaricare nulla e contatta subito il responsabile della beta.", LIGHT_RED)

    doc.add_heading("12. Risoluzione dei problemi", level=1)
    issues = [
        ("Non riesco ad accedere", "Controlla tenant ID, email e password. Evita spazi iniziali/finali. Chiedi un reset al responsabile; non creare un nuovo tenant."),
        ("Il file non viene accettato", "Verifica estensione CSV/XLSX/XLS, dimensione, integrità e presenza di una tabella con intestazioni."),
        ("L'analisi resta in elaborazione", "Attendi alcuni minuti e aggiorna una volta. Annota l'ID. Se persiste, contatta il supporto senza inviare il file via canali non autorizzati."),
        ("L'analisi fallisce", "Leggi il messaggio, semplifica la domanda, controlla intestazioni e tipi delle colonne, quindi riprova con un file ridotto."),
        ("Il risultato usa la colonna sbagliata", "Segna Parziale o Errato e indica nella nota la colonna richiesta e quella utilizzata."),
        ("Il totale non coincide", "Controlla duplicati, null, filtri, righe di totale nel file e definizione dell'unità di conteggio. Salva il feedback con il valore atteso."),
        ("Non vedo il feedback", "Il modulo compare nelle analisi completate. Torna allo storico, riapri il risultato e scorri fino in fondo."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [3000, 6360])
    for i, text in enumerate(("Problema", "Cosa fare")):
        shade(table.rows[0].cells[i], "DCE6F1")
        set_font(table.rows[0].cells[i].paragraphs[0].add_run(text), 10, True, NAVY)
    for problem, action in issues:
        cells = table.add_row().cells
        set_font(cells[0].paragraphs[0].add_run(problem), 9.5, True, NAVY)
        set_font(cells[1].paragraphs[0].add_run(action), 9.5, False, NAVY)
    set_table_geometry(table, [3000, 6360])

    doc.add_heading("13. Checklist della sessione beta", level=1)
    doc.add_paragraph("Completa questa sequenza per ogni prova significativa.")
    checks = [
        "Ho usato un file autorizzato e comprensibile.",
        "Ho annotato la domanda esatta e il risultato atteso.",
        "Ho controllato righe, colonne e tipi di dato.",
        "Ho avviato una sola analisi e atteso lo stato finale.",
        "Ho verificato che il report risponda alla domanda originale.",
        "Ho ricalcolato almeno il numero principale in modo indipendente.",
        "Ho scelto Corretto, Parziale o Errato secondo i criteri del manuale.",
        "Ho assegnato un voto coerente e scritto una nota concreta.",
        "Ho salvato il feedback.",
        "Ho comunicato eventuali problemi critici al responsabile.",
    ]
    for item in checks:
        bullet(doc, "[ ] " + item)

    doc.add_heading("Glossario essenziale", level=1)
    glossary = [
        ("Analisi", "Una singola richiesta composta da domanda, file, elaborazione e risultato."),
        ("CSV", "File di testo tabellare in cui i valori sono separati da virgole o altri delimitatori."),
        ("Excel", "Cartella di lavoro .xlsx o .xls contenente una tabella dati."),
        ("Dataset", "L'insieme di righe e colonne caricato per l'analisi."),
        ("Deterministico", "Calcolo ripetibile basato sui dati e su regole esplicite, senza inventare valori."),
        ("Feedback", "Valutazione salvata dall'utente sul risultato ricevuto."),
        ("Knowledge Graph", "Mappa delle relazioni tra analisi, dataset, colonne, insight e report del tenant."),
        ("Narrativa opzionale", "Riformulazione testuale destinata alla leggibilità; non sostituisce il calcolo autorevole."),
        ("Tenant", "Spazio isolato assegnato a una singola organizzazione."),
    ]
    for term, definition in glossary:
        p = doc.add_paragraph()
        set_font(p.add_run(term + ": "), 10.5, True, TEAL)
        set_font(p.add_run(definition), 10.5, False, NAVY)

    doc.add_heading("Scheda da inviare al supporto", level=1)
    doc.add_paragraph("Se devi segnalare un problema, invia queste informazioni senza allegare dati riservati salvo autorizzazione esplicita:")
    for text in [
        "Data e ora della prova",
        "ID analisi visibile nello storico",
        "Stato finale",
        "Domanda esatta",
        "Tipo di file e numero approssimativo di righe/colonne",
        "Comportamento osservato e comportamento atteso",
        "Screenshot privo di dati sensibili, se consentito",
    ]:
        bullet(doc, text)
    callout(doc, "FINE DELLA PROVA", "Dopo aver salvato il feedback, torna al portale e premi Esci. Grazie: una valutazione accurata è più utile di molte prove eseguite senza verifica.", LIGHT_GREEN)


def prepare_images():
    required = [
        "01_accesso_registrazione.png", "02_home_nuova_analisi.png",
        "03_analisi_in_elaborazione.png", "04_storico_completato.png",
        "05_risultato_top.png", "05_feedback.png", "06_knowledge_graph.png",
    ]
    missing = [name for name in required if not (ASSETS / name).exists()]
    if missing:
        raise FileNotFoundError("Immagini mancanti: " + ", ".join(missing))


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    prepare_images()
    doc = Document()
    configure(doc)
    add_cover(doc)
    add_manual(doc)
    core = doc.core_properties
    core.title = "Manuale utente illustrato Veraxis - Private Beta"
    core.subject = "Guida completa per utenti senza esperienza precedente"
    core.author = "Veraxis"
    core.keywords = "Veraxis, private beta, manuale utente, feedback"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
