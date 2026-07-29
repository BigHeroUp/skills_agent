from io import BytesIO

from docx import Document

from scripts.build_internal_validation_campaign import execute_campaign, build_docx
from scripts.import_internal_validation_feedback import normalize_outcome, parse_score, summarize


def test_internal_campaign_has_20_independent_review_records():
    campaign = execute_campaign()
    assert campaign["total"] == 20
    assert campaign["automatic_passed"] == 20
    assert len({item["review_id"] for item in campaign["cases"]}) == 20
    assert all(item["human_feedback"]["outcome"] is None for item in campaign["cases"])


def test_campaign_docx_contains_feedback_form_for_every_case(tmp_path):
    output = tmp_path / "campaign.docx"
    build_docx(execute_campaign(), output)
    document = Document(output)
    feedback_tables = [table for table in document.tables if table.cell(0, 0).text == "Outcome"]
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert len(feedback_tables) == 20
    assert "REV-001" in text and "REV-020" in text
    assert all("[COMPILARE" in table.cell(0, 1).text for table in feedback_tables)


def test_feedback_import_validation_and_summary_contract():
    assert normalize_outcome("Corretto") == "correct"
    assert normalize_outcome("parziale") == "partial"
    assert parse_score("5") == 5
    assert parse_score("0") is None
    cases = [
        {"human_feedback": {"outcome": "correct", "rating": 5}},
        {"human_feedback": {"outcome": "partial", "rating": 3}},
    ]
    assert summarize(cases) == {
        "total": 2,
        "outcomes": {"correct": 1, "partial": 1, "incorrect": 0},
        "average_rating": 4.0,
    }
